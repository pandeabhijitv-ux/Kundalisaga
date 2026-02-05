"""
Document Processor Module
Handles ingestion of astrology books in various formats
"""
import os
import hashlib
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass
import pdfplumber
from docx import Document as DocxDocument
import pytesseract
from PIL import Image
from tqdm import tqdm

from src.utils import logger, config


@dataclass
class ProcessedDocument:
    """Represents a processed document"""
    file_path: str
    file_name: str
    file_hash: str
    text_content: str
    page_count: int
    file_type: str
    metadata: Dict


class DocumentProcessor:
    """Process various document formats and extract text"""
    
    def __init__(self):
        self.logger = logger
        self.supported_formats = config.get('documents.supported_formats', [])
    
    def process_file(self, file_path: str) -> Optional[ProcessedDocument]:
        """
        Process a single file and extract text
        
        Args:
            file_path: Path to the file
        
        Returns:
            ProcessedDocument or None if failed
        """
        path = Path(file_path)
        
        if not path.exists():
            self.logger.error(f"File not found: {file_path}")
            return None
        
        file_ext = path.suffix.lower()
        
        if file_ext not in self.supported_formats:
            self.logger.warning(f"Unsupported format: {file_ext}")
            return None
        
        self.logger.info(f"Processing: {path.name}")
        
        try:
            # Extract text based on file type
            if file_ext == '.pdf':
                text, page_count = self._process_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text, page_count = self._process_docx(file_path)
            elif file_ext == '.txt':
                text, page_count = self._process_txt(file_path)
            elif file_ext in ['.png', '.jpg', '.jpeg']:
                text, page_count = self._process_image(file_path)
            else:
                return None
            
            # Calculate file hash for deduplication
            file_hash = self._calculate_hash(file_path)
            
            # Create metadata
            metadata = {
                'file_size': path.stat().st_size,
                'created_at': path.stat().st_ctime,
                'modified_at': path.stat().st_mtime,
            }
            
            return ProcessedDocument(
                file_path=str(path),
                file_name=path.name,
                file_hash=file_hash,
                text_content=text,
                page_count=page_count,
                file_type=file_ext,
                metadata=metadata
            )
        
        except Exception as e:
            self.logger.error(f"Error processing {path.name}: {str(e)}")
            return None
    
    def _process_pdf(self, file_path: str) -> tuple:
        """Extract text from PDF with encoding handling"""
        text_content = []
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    # Clean the text to remove garbled/corrupted content
                    cleaned_text = self._clean_extracted_text(text)
                    if cleaned_text:
                        text_content.append(cleaned_text)
        
        combined_text = '\n\n'.join(text_content)
        
        # If most of the text is garbled, warn and return minimal content
        if self._is_mostly_garbled(combined_text):
            self.logger.warning(f"PDF {file_path} contains heavily corrupted/non-English text. Consider providing English translated versions for better results.")
            return combined_text[:500] + "\n\n[Note: This PDF contains text encoding issues. For better results, please provide English language astrology books or properly encoded PDFs.]", page_count
        
        return combined_text, page_count
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean extracted text by removing garbled content"""
        import re
        
        # Remove excessive special characters and symbols
        # Keep basic punctuation and spaces
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Skip lines with excessive special characters
            special_char_count = len(re.findall(r'[^a-zA-Z0-9\s.,;:!?()\-\'\"]', line))
            total_chars = len(line.strip())
            
            if total_chars == 0:
                continue
                
            # If more than 40% special characters, skip this line
            if total_chars > 0 and (special_char_count / total_chars) > 0.4:
                continue
            
            # Keep lines that have meaningful English content
            word_chars = len(re.findall(r'[a-zA-Z]', line))
            if word_chars > 10 or total_chars < 20:  # Either has English words or is short
                cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _is_mostly_garbled(self, text: str) -> bool:
        """Check if text is mostly garbled/corrupted"""
        if not text or len(text) < 100:
            return True
        
        import re
        # Count readable English characters vs special characters
        english_chars = len(re.findall(r'[a-zA-Z]', text))
        total_chars = len(text.replace('\n', '').replace(' ', ''))
        
        if total_chars == 0:
            return True
        
        # If less than 30% English characters, consider it garbled
        return (english_chars / total_chars) < 0.3
    
    def _process_docx(self, file_path: str) -> tuple:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text_content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)
        
        return '\n\n'.join(text_content), len(doc.paragraphs)
    
    def _process_txt(self, file_path: str) -> tuple:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        return text, text.count('\n') + 1
    
    def _process_image(self, file_path: str) -> tuple:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            return text, 1
        except Exception as e:
            self.logger.warning(f"OCR failed for {file_path}: {str(e)}")
            return "", 1
    
    def _calculate_hash(self, file_path: str) -> str:
        """Calculate SHA256 hash of file"""
        sha256_hash = hashlib.sha256()
        
        with open(file_path, "rb") as f:
            for byte_block in iter(lambda: f.read(4096), b""):
                sha256_hash.update(byte_block)
        
        return sha256_hash.hexdigest()
    
    def process_directory(self, directory_path: str) -> List[ProcessedDocument]:
        """
        Process all supported files in a directory
        
        Args:
            directory_path: Path to directory
        
        Returns:
            List of ProcessedDocument objects
        """
        path = Path(directory_path)
        
        if not path.exists() or not path.is_dir():
            self.logger.error(f"Invalid directory: {directory_path}")
            return []
        
        # Find all supported files
        all_files = []
        for ext in self.supported_formats:
            all_files.extend(path.rglob(f"*{ext}"))
        
        self.logger.info(f"Found {len(all_files)} files to process")
        
        # Process files
        processed_docs = []
        
        for file_path in tqdm(all_files, desc="Processing documents"):
            doc = self.process_file(str(file_path))
            if doc and doc.text_content.strip():
                processed_docs.append(doc)
        
        self.logger.info(f"Successfully processed {len(processed_docs)} documents")
        return processed_docs
    
    def chunk_text(self, text: str, chunk_size: int = 1000, 
                   chunk_overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to split
            chunk_size: Size of each chunk
            chunk_overlap: Overlap between chunks
        
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < text_length:
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > chunk_size * 0.5:  # At least 50% of chunk
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - chunk_overlap
        
        return chunks
